from pydub import AudioSegment
import math
import os
import openai
from openai import OpenAI
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime  # Importera endast datetime-funktionen
import re
import shutil
from openai import OpenAI
import sys 
import logging
import configparser
import csv
import re
from main import extract_timestamp_from_filename

config = configparser.ConfigParser()
config.read('config.ini')

logging.basicConfig(level=logging.INFO,
                    filename='myapp.log',
                    filemode='w',
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                    datefmt='%Y-%m-%d %H:%M:%S')
logging.info("transkript.py startas")
print(f"transkript.py startas")
# OpenAI API-nycklar
openai.api_key = open("key.txt", "r").read().strip('\n')
client = OpenAI(api_key=openai.api_key)

# Sökvägar
input_mp3_directory = "E:\\Dropbox\\Meetings\\Audio"
output_word_directory = "E:\\Dropbox\\Meetings\\Transcriptions"
processed_mp3_directory = "E:\\Dropbox\\Meetings\\Audio\\Transcribed_Mp3"


# Säkerställ att katalogerna existerar
os.makedirs(output_word_directory, exist_ok=True)
os.makedirs(processed_mp3_directory, exist_ok=True)

def extract_simple_timestamp(filename):
    # Ta bort filändelsen och dela upp resten av strängen
    name_without_extension = filename.split('.')[0]
    parts = name_without_extension.split('_')

    # Kontrollera att vi har tillräckligt många delar för att innehålla datum och tid
    if len(parts) >= 3:
        # Datumet är det tredje sista elementet och tiden det andra sista
        date_part = parts[-2]  # Tredje sista delen är datumdelen
        time_part = parts[-1]  # Andra sista delen är tidsdelen

        # Sätt ihop och returnera som en sträng
        return f"{date_part}_{time_part}"
    else:
        print(f"Kunde inte extrahera tidsstämpel från filnamnet '{filename}'")
        return None


def get_company_name(selected_project):
    logging.info(f"Söker företagsnamn för projekt: '{selected_project}'")
    try:
        with open('E:\\Dropbox\\CODE\\Transkript\\venvTranskript\\projects.csv', mode='r', encoding='utf-8') as file:
            reader = csv.DictReader(file, delimiter=';')  # Använd semikolon som avgränsare
            for row in reader:
                logging.info(f"Läser rad: {row}")
                if row['project_name'].strip() == selected_project.strip():
                    logging.info(f"Hittade matchning: {row['company_name']}")
                    return row['company_name']
            logging.warning(f"Inget matchande projekt hittades för: '{selected_project}'")
            return "Okänt Företag"
    except FileNotFoundError as e:
        logging.error(f"CSV-filen hittades inte: {e}")
        raise
    except KeyError as e:
        logging.error(f"Felaktigt fältnamn i CSV-filen: {e}")
        raise
    except Exception as e:
        logging.error(f"Oväntat fel vid läsning av CSV-filen: {e}")
        raise

# Din befintliga split_mp3-funktion här...
def split_mp3(file_path, segment_size_mb=25):
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file {file_path} does not exist.")

        # Calculate the file size in MB
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)

        # If the file size is smaller than the segment size, no splitting is needed
        if file_size_mb <= segment_size_mb:
            print(f"The file is smaller than {segment_size_mb}MB, no segmentation needed.")
            return [file_path]

        # Load the audio file
        audio = AudioSegment.from_mp3(file_path)

        # Calculate the total duration in milliseconds
        total_duration_ms = len(audio)

        # Calculate the duration of each segment in milliseconds
        # We assume the bit rate of the mp3 is 128kbps for calculation
        segment_duration_ms = (segment_size_mb * 1024 * 8) / 128 * 1000

        # Calculate the number of segments needed
        num_segments = math.ceil(total_duration_ms / segment_duration_ms)

        # Split and export the segments
        segment_paths = []
        for i in range(num_segments):
                start_ms = i * segment_duration_ms
                end_ms = min((i + 1) * segment_duration_ms, total_duration_ms)
                segment = audio[start_ms:end_ms]
                segment_path = f"{file_path}_segment_{i + 1}.mp3"
                segment.export(segment_path, format="mp3")
                segment_paths.append(segment_path)
                print(f"Segment {i + 1} exported as {segment_path}.")
        print("Segmentering av filer klar")
        return segment_paths
    except FileNotFoundError as e:
        print(f"Filfel: {e}")
        return []  # Returnera en tom lista om filen inte finns
    except Exception as e:
        print(f"Oväntat fel vid segmentering av MP3: {e}")
        return []
    
def transcribe_audio_list(segments):
    combined_transcription = ""
    for audio_file_path in segments:
        try:
            with open(audio_file_path, 'rb') as audio_file:
                transcription = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio_file)
                combined_transcription += transcription.text + " "
            print("Transkription av kombinerade segment klar")
        except FileNotFoundError as e:
            print(f"Filfel: {e}")
        except Exception as e:
            print(f"Oväntat fel vid transkribering: {e}")
    return combined_transcription
        
# Din befintliga abstract_summary_extraction-funktion här...
def abstract_summary_extraction(transcription):
    try:
        response = client.chat.completions.create(
            model="gpt-4-1106-preview",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "Du är en högt kvalificerad AI utbildad i språkförståelse och sammanfattning. Du har mycket hög kunskap om IT-relaterade frågor. Jag skulle vilja att du läser följande transkription av ett möte och sammanfattar det till ett koncist abstrakt stycke. Sikta på att behålla de viktigaste punkterna, och ge en sammanhängande och läsbar sammanfattning som kan hjälpa en person att förstå huvudpunkterna i diskussionen utan att behöva läsa hela texten. Var vänlig och undvik onödiga detaljer eller tangentiella punkter."
                },
                {
                    "role": "user",
                    "content": transcription
                }
            ]
        )
        # return response['choices'][0]['message']['content']  # Format from old API
        response = response.choices[0].message.content
        print("Abstract summary klar")
        return response
    except Exception as e:
        print(f"Oväntat fel vid sammanfattning: {e}")
        return "" # Returnera en tom sträng om något går fel

# Din befintliga key_points_extraction-funktion här...
def key_points_extraction(transcription):
    try:
        response = client.chat.completions.create(
            model="gpt-4-1106-preview",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "Du är en skicklig AI med en specialitet i att sammanfatta information till viktiga punkter. Du har mycket hög kunskap om IT-relaterade frågor. Baserat på följande text, identifiera och lista de huvudpunkter som diskuterades eller togs upp. Dessa bör vara de mest betydelsefulla idéerna, fynden, slutsaterna eller ämnena som är avgörande för diskussionens kärna. Ditt mål är att ge en lista som någon kan läsa för att snabbt förstå vad som diskuterades."
                },
                {
                    "role": "user",
                    "content": transcription
                }
            ]
        )
        response = response.choices[0].message.content
        print("Keypoints klara")
        return response
    except Exception as e:
        print(f"Oväntat fel vid key points: {e}")
        return "" # Returnera en tom sträng om något går fel    
        
# Din befintliga action_item_extraction-funktion här...
def action_item_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4-1106-preview",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "Du är en AI-expert på att analysera samtal och extrahera åtgärdspunkter. Du har mycket hög kunskap om IT-relaterade frågor. Detta möte kan innehålla tekniska termer från IT-branschen, ta hänsyn till detta. Var vänlig och granska texten och identifiera eventuella uppgifter, uppdrag eller åtgärder som avtalades eller nämndes som något som behöver göras. Detta kan vara uppgifter tilldelade specifika individer, eller allmänna åtgärder som gruppen har beslutat att vidta. Var vänlig och lista dessa åtgärdspunkter på ett klart och koncist sätt."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    response = response.choices[0].message.content
    print("Actions klara")
    return response

# Din befintliga extract_mentioned_non_participants-funktion här...
def extract_mentioned_non_participants(transcription, participants):
    participants_str = ', '.join(participants)
    prompt = (
        "Du är en AI-expert på att analysera samtal och extrahera namn och roller på de talande personerna. Var vänlig och granska texten och identifiera varje person som nämns i diskussionerna, deras titel eller roll, och annan personlig information de tillhandahåller såsom plats. Se till att granska hela konversationen och inkludera nya personer som nämns senare i mötet. Mötet kan vara ett företagsrapporteringssamtal med analytiker; om så är fallet, se till att inkludera analytikerna som ställer frågor senare i samtalet. Var vänlig och lista alla mötesdeltagare och deras relaterade information på ett klart och koncist sätt. Om det finns tydliga grupper av personer, såsom kunder, konnsulter, leverantörer eller annat gruppera dem därefter"
        "Identifiera alla personer som omnämns i följande text, "
        "men som inte är listade som deltagare i mötet. "
        "Mötesdeltagare inkluderar: " + participants_str + ".\n\n"
        "Transkription:\n" + transcription
    )

    response = client.chat.completions.create(
        model="gpt-4-1106-preview",
        temperature=0.5,
        messages=[
            {
                "role": "system",
                "content": prompt
            }
        ]
    )
    mentioned_non_participants = response.choices[0].message.content
    print("mentioned_non_participants klar")
    return mentioned_non_participants

# Din befintliga participant_list-funktion här...
def participant_list(transcription):
    response = client.chat.completions.create(
        model="gpt-4-1106-preview",
        temperature=0,
        messages=[{
            "role": "system",
            "content": "..."  # Din befintliga prompt här
        }]
    )
    response_text = response.choices[0].message.content

    # Antag att vi kan extrahera deltagarnamn och roller härifrån
    # Om vi inte kan, generera "Deltagare X"-namn
    extracted_participants = extract_participants(response_text)
    if not extracted_participants:
        extracted_participants = generate_anonymous_participants(transcription)
    return extracted_participants

def extract_participants(response_text):
    # Implementera logik för att extrahera deltagarnamn och roller
    # Returnera en lista med namn och roller
    pass
def generate_anonymous_participants(transcription):
    num_participants = count_participants(transcription)
    # Om num_participants är None eller 0, sätt det till ett standardvärde
    num_participants = num_participants if num_participants else 1
    return [f"Deltagare {i+1}" for i in range(num_participants)]

def count_participants(transcription):
    # Implementera en metod för att räkna antal unika deltagare
    # Här är en förenklad exempelimplementation:
    # Anta att varje ny rad i transkriptionen indikerar en ny deltagare
    # Detta är naturligtvis inte alltid korrekt, men fungerar som en nödlösning
    return len(transcription.split('\n')) if transcription else 0


# Din befintliga ioi_extraction-funktion här...
def ioi_extraction(transcription):  # Items of interest
    with open('ioi.txt', 'r', encoding='utf-8') as file:  # Read items of interest from file
        ioi = file.read()
    ioi = ioi.replace('\n', ', ').strip(', ')  # Replace line breaks with commas
    content = "Du är en AI-expert på att analysera företagsmöten inom IT-sektorn, och extrahera nyckelobjekt av intresse som specificerats av användaren. Var vänlig och granska hela texten noggrant och identifiera om någon av följande termer nämns i transkriptionen av detta möte. Börja med att dela upp transkriptionen i mindre delar på mindre än 10 000 tecken vardera. Sök noggrant igenom varje del efter de angivna nyckeltermerna. Om någon av termerna nämns, gör två saker: 1) upprepa exakt vad som sades om den termen, och 2) förklara vad som menades med diskussionen om den termen. Presentera resultatet på ett organiserat sätt. När du är klar, granska transkriptionen igen för att säkerställa att inga av de angivna termerna missades. Här är listan över termer: "
    content += ioi
    response = client.chat.completions.create(
        model="gpt-4-1106-preview",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": content
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    response = response.choices[0].message.content
    print("IOI-extraction klar")
    return response


# Din befintliga sentiment_analysis-funktion här...
def sentiment_analysis(transcription):
    response = client.chat.completions.create(
        model="gpt-4-1106-preview",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "Som en AI med expertis inom språk- och känslanalys är din uppgift att analysera känsloläget i följande text. Vänligen beakta den övergripande tonen i diskussionen, känslorna som förmedlas genom det använda språket och sammanhanget där ord och fraser används. Ange om känsloläget är övervägande positivt, negativt eller neutralt och ge korta förklaringar till din analys där det är möjligt."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
def create_word_document(transcription, summary, key_points, action_items, 
                         ioi_discussion, sentiment, segment_time, transcribe_time, summary_time, key_points_time, 
                         action_items_time, participants_time, ioi_time, sentiment_time, total_time, title, participants_intro, 
                         mentioned_non_participants, doc_file_name, selected_project, meeting_type, timestamp, word_path):
    company_name = get_company_name(selected_project)
    doc_title = f"{company_name} - {selected_project}"

    try:
        doc = Document()
        # Titel
        doc.add_heading(doc_title, level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Mötestyp
        doc.add_heading(meeting_type, level=2).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Datum
        doc.add_heading(timestamp, level=3).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        # 1. Närvarande
        doc.add_heading('1. Närvarande', level=2)
        for participant in participants_intro:
            doc.add_paragraph(participant)  # Se till att varje deltagare är på en ny rad
            print(f"Deltagare: {participant}")
        # 2. Summering
        doc.add_heading('2. Summering', level=2)
        doc.add_paragraph(summary)

        # 3. Huvudpunkter
        doc.add_heading('3. Huvudpunkter', level=2)
        doc.add_paragraph(key_points)

        # 4. Åtgärder
        doc.add_heading('4. Åtgärder', level=2)
        doc.add_paragraph(action_items)

        # 5. Specifika Ämnen
        doc.add_heading('5. Specifika Ämnen', level=2)
        doc.add_paragraph(ioi_discussion)
    
        # 6. Omnämnda Personer
        doc.add_heading('6. Omnämnda Personer', level=2)
        for person in mentioned_non_participants.split('\n'):
            doc.add_paragraph(person)
        
        # 7. Stämningsanalys
        doc.add_heading('7. Stämningsanalys', level=2)
        doc.add_paragraph(sentiment)

        # 8. Transkribering
        doc.add_heading('8. Transkribering', level=2)
        doc.add_paragraph(transcription)


        full_file_path = os.path.join(word_path, doc_file_name)
        doc.save(full_file_path)
        print(f"Dokument sparad som {full_file_path}")
    except Exception as e:
        print(f"Ett oväntat fel inträffade då Word-dokumentet skapades: {e}")
def load_corrections(file_path):
    corrections = {}
    with open(file_path, 'r', encoding='utf-8') as file:
        for line in file:
            parts = line.strip().split(', ')
            if len(parts) == 2:
                corrections[parts[0]] = parts[1]
    print("Load corrections klar")
    return corrections

def apply_corrections(text, corrections):
    for wrong, right in corrections.items():
        text = text.replace(wrong, right)
    print("Rättelser klara")
    return text

def process_transcription():
    # Läs valen från filen
    with open("selected_options.txt", "r") as file:
        selected_project = file.readline().strip()
        company_name = file.readline().strip()
        meeting_type = file.readline().strip()
        participants_intro = file.readline().strip().split(", ")

    # timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M')
    transcription = summary = key_points = action_items = ""
    ioi_discussion = sentiment = title = mentioned_non_participants = ""
    segment_time = transcribe_time = summary_time = key_points_time = action_items_time = participants_time = ioi_time = sentiment_time = total_time = 0

    mp3_files = [file for file in os.listdir(input_mp3_directory) if file.endswith('.mp3')]
    if not mp3_files:
        print("Inga MP3-filer hittades i katalogen:", input_mp3_directory)
        return

    for file_name in mp3_files:
        file_path = os.path.join(input_mp3_directory, file_name)
        # Extrahera bara filnamnet
        filename_only = os.path.basename(file_path)
        print(f"Filnamn: {filename_only}")
        # Använd funktionen för att extrahera och formatera tidsstämpeln
        timestamp = extract_simple_timestamp(filename_only)
        if timestamp:
            print("Extraherad tidsstämpel:", timestamp)
        else:
            print("Kunde inte extrahera tidsstämpel från filnamnet")


        # Bearbetningen av varje fil börjar här
        t0 = time.time()
        segments = split_mp3(file_path)
        t1 = time.time()

        transcription = transcribe_audio_list(segments)
        t2 = time.time()

        correction_map = load_corrections("corrections.txt")
        transcription = apply_corrections(transcription, correction_map)

        t3 = time.time()
        summary = abstract_summary_extraction(transcription)
        t4 = time.time()
        key_points = key_points_extraction(transcription)
        t5 = time.time()
        action_items = action_item_extraction(transcription)
        t6 = time.time()

        participants = participant_list(transcription)
        mentioned_non_participants = extract_mentioned_non_participants(transcription, participants_intro)
        t7 = time.time()

        ioi_discussion = ioi_extraction(transcription)
        t8 = time.time()

        sentiment = sentiment_analysis(transcription)
        t9 = time.time()

        t10 = time.time()

        segment_time = t1 - t0
        transcribe_time = t2 - t1
        summary_time = t4 - t3
        key_points_time = t5 - t4
        action_items_time = t6 - t5
        participants_time = t7 - t6
        ioi_time = t8 - t7
        sentiment_time = t9 - t8
        total_time = t10 - t0

     # Skapa Word-dokument
    new_filename_base = f"{selected_project}_{meeting_type}_{timestamp}"
    doc_file_name = f"{new_filename_base}.docx"
    print("Typ av participants_intro:", type(participants_intro))
    print("Innehåll i participants_intro:", participants_intro)

    create_word_document(
    transcription, summary, key_points, action_items, 
    ioi_discussion, sentiment, 
    segment_time, transcribe_time, summary_time, key_points_time, 
    action_items_time, participants_time, ioi_time, sentiment_time, total_time, 
    title, participants_intro, mentioned_non_participants, 
    doc_file_name, selected_project, meeting_type, timestamp, word_path
)

        
    # Flytta den bearbetade MP3-filen
    processed_file_path = os.path.join(processed_mp3_directory, file_name)
    shutil.move(file_path, processed_file_path)
    print(f"Fil flyttad till: {processed_file_path}")

    # Radera segmentfilerna
    for segment_path in segments:
        if os.path.exists(segment_path):
            os.remove(segment_path)
            print(f"Segmentfil raderad: {segment_path}")
        else:
            print(f"Kunde inte hitta filen att radera: {segment_path}")

def main():
    try:
        process_transcription()
    except Exception as e:
        print(f"Oväntat fel: {e}")

if __name__ == "__main__":
    main()